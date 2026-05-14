"""Rewrite the team proposal docx with NLP injection canary content.

Strategy:
- Open the original docx
- Update course header text on paragraphs [15], [16], and project title on [33]
- Replace body content between section headings, preserving heading styles
- Save to the original path (the backup is kept under .original.bak.docx)
"""

from copy import deepcopy
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "/path/to/file"
DST = SRC  # overwrite

doc = docx.Document(SRC)


def set_paragraph_text(p, text, bold=None):
    """Replace all runs in a paragraph with a single run carrying text."""
    # Remove existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    if bold is not None:
        run.bold = bold
    return run


def insert_paragraph_after(paragraph, text="", style=None, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Insert a new paragraph immediately after the given one. Returns the new paragraph."""
    new_p = deepcopy(paragraph._p)
    # Clear children of the new paragraph except pPr
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if alignment is not None:
        new_para.alignment = alignment
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def delete_paragraph(paragraph):
    p = paragraph._p
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


paragraphs = list(doc.paragraphs)

# 1) Course header text
set_paragraph_text(paragraphs[15], "NATURAL LANGUAGE PROCESSING", bold=True)
set_paragraph_text(paragraphs[16], "Spring 25-26", bold=True)

# 2) Project title
new_title = "“Prompt Injection Canary: A Two-Pass Detector that Extracts Inferred Adversarial Instructions”"
set_paragraph_text(paragraphs[33], new_title, bold=True)

# 3) Body sections.
# Headings stay at indexes [35, 38, 45, 58, 64]. Everything in between is body.
# Strategy: keep heading paragraphs, delete current body paragraphs in each section,
# then insert new body content after the heading.

# Identify section heading paragraph objects
headings = {
    "intro": paragraphs[35],
    "problem": paragraphs[38],
    "solution": paragraphs[45],
    "testing": paragraphs[58],
    "references": paragraphs[64],
}

# To safely delete body paragraphs, identify them by index ranges in the *original* list,
# then delete from the end forward so indexes do not shift.
body_ranges = [
    ("intro", 36, 37),
    ("problem", 39, 44),
    ("solution", 46, 57),
    ("testing", 59, 63),
    ("references", 65, 69),
]

# Refresh paragraphs reference for deletion safety
to_delete = []
for _, start, end in body_ranges:
    for i in range(start, end + 1):
        to_delete.append(paragraphs[i])
# Delete in reverse to keep XML traversal clean
for p in reversed(to_delete):
    if p._p is not None:
        delete_paragraph(p)

# Now insert new body content after each heading. We insert in reverse order per section
# (insert_after places the new paragraph as the next sibling, so to keep order we insert
# each paragraph after the last inserted one).

def append_body(after_para, lines):
    """lines is a list of (text, bold) tuples. Returns the last inserted paragraph."""
    current = after_para
    for text, bold in lines:
        current = insert_paragraph_after(current, text=text, style="Normal", bold=bold,
                                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    return current


# ----- Section 1, Introduction and Background -----
intro_lines = [
    ("Large language model (LLM) applications increasingly consume untrusted text. Browser agents read web pages, retrieval pipelines pull documents, MCP servers return tool outputs, and chat assistants accept user uploads. Any of that text can carry a prompt injection, where embedded natural language instructions hijack the host application's behavior. Prompt injection is the OWASP LLM01 risk for 2025 and the most cited emerging threat vector across security vendor reports from 2025 onward.", False),
    ("Public detectors approach the problem as binary or three-class classification. Meta's PromptGuard, ProtectAI's deberta-v3-prompt-injection, Deepset's deberta-v3-injection, and Lakera Guard all emit a single label. Recent work shows two material gaps. First, single-pass detectors collapse under paraphrase, translation round-trip, and non-English rewrites, with published evasion success rates approaching 100 percent. Second, a binary label is the wrong downstream API. A host agent that learns its input is injected still does not know what the attacker was trying to make it do, so it cannot refuse precisely or escalate effectively.", False),
]
last = append_body(headings["intro"], intro_lines)

# ----- Section 2, Problem Description and Scope -----
problem_lines = [
    ("This project proposes Prompt Injection Canary, a two-pass detector that flags adversarial inputs and extracts the inferred adversarial instruction in canonical form. The scope covers direct injections (in the user message), indirect injections (in retrieved documents and tool outputs), and three robustness axes that current detectors are documented to fail on: paraphrase, multilingual rewrites (English to Turkish, Russian, and Chinese and back), and encoded payloads (base64, ROT13, zero-width characters).", False),
    ("The project aims to:", False),
    ("Build a two-pass detector that returns a structured object (extracted instruction, attacker-goal tag, confidence) rather than a single label.", False),
    ("Define and measure a new metric, extraction faithfulness, that captures how well the extracted instruction matches the ground-truth attacker goal.", False),
    ("Release an open robustness benchmark with paraphrase, translation, and encoded splits, usable against any detector that exposes a predict(text) interface.", False),
]
last = append_body(headings["problem"], problem_lines)

# ----- Section 3, Proposed Solution -----
solution_lines = [
    ("To achieve the project's objectives, the following solution is proposed:", False),
    ("Data Curation:", True),
    ("Train data combines HackAPrompt (over 600,000 adversarial prompts), Tensor Trust (563,000 attacks and 118,000 defenses), BIPIA train split, and InjecAgent train split. Benign distribution comes from Dolly 15k, Alpaca, and a filtered ShareGPT sample. Each injection is augmented with one LLM paraphrase, two translation round-trips (English to Turkish and Russian), and one encoding variant. This yields roughly a five-times expansion of the injection class.", False),
    ("Pass 1, Fast Multilingual Classifier:", True),
    ("An mDeBERTa-v3-base encoder, warm-started from Meta's PromptGuard-2-86M, fine-tuned with a contrastive objective on (clean, injected, paraphrased, translated) tuples. Outputs an injection probability that is thresholded against a calibrated decision boundary set on a held-out validation set for a target false-positive rate of 1 percent.", False),
    ("Pass 2, Instruction Extractor:", True),
    ("An instruction-tuned LLM (Qwen 2.5 7B Instruct) is invoked only when Pass 1 flags an input. Using JSON-mode constrained decoding, it produces: (i) the smallest span of the input that contains adversarial instructions, (ii) the instruction restated in canonical imperative form, (iii) an attacker-goal tag from a five-class taxonomy (exfiltration, hijack, refusal override, tool misuse, social engineering), and (iv) a confidence value in [0, 1]. The structured output is the canary string the host application uses to refuse precisely or escalate to a human reviewer.", False),
    ("Baselines and Reference Implementation:", True),
    ("Seven baselines are benchmarked: a random/majority floor, TF-IDF plus logistic regression, Deepset deberta-v3-base-injection, ProtectAI deberta-v3-base-prompt-injection v2, Meta PromptGuard-2-86M, Attention Tracker on Llama 3.1 8B, and GPT-4o-mini as a zero-shot judge. The full pipeline is released as a reproducible GitHub repository with a Dockerfile and a HuggingFace dataset card for the robustness benchmark.", False),
]
last = append_body(headings["solution"], solution_lines)

# ----- Section 4, Plan for Testing and Result Analysis -----
testing_lines = [
    ("Evaluation Splits:", True),
    ("The test set is partitioned into six disjoint splits, each held out from training: Direct-EN (HackAPrompt and Tensor Trust, around 1,500 items), Indirect-EN (BIPIA test split, 1,250 items), Paraphrased (1,000 items generated by an unseen LLM paraphraser), Translated TR/RU/ZH (1,200 items combined), Encoded (600 items in base64, ROT13, and zero-width form), and Overdefense (800 benign prompts containing trigger words such as “ignore” or “system,” constructed in the InjecGuard style).", False),
    ("Performance Metrics:", True),
    ("For classification: macro F1, AUROC, and false-positive rate at a fixed true-positive rate of 0.95. For extraction quality: exact match against the five-goal taxonomy, span F1 against a manually labeled gold set of 300 items, and LLM-as-judge agreement using Claude 3.5 Sonnet on a 500-sample subset. For deployability: end-to-end latency at p50 and p95, and cost per 1,000 inputs at public API pricing.", False),
    ("Validation Techniques:", True),
    ("Bootstrap 95 percent confidence intervals are reported for every headline number. McNemar’s test is used for pairwise classifier comparisons. Calibration plots and expected calibration error are reported for any model whose confidence value is consumed downstream. Ablations include Pass 1 only versus the full canary, Pass 2 only versus the cascade, leave-one-out by augmentation type, and taxonomy granularity (five tags versus three tags versus binary).", False),
    ("Results will be documented clearly with figures and tables, interpreted with attention to error modes (especially where paraphrase or translation breaks each baseline), and packaged into a 2 to 6 page academic report alongside the code repository and the final presentation.", False),
]
last = append_body(headings["testing"], testing_lines)

# ----- Section 5, References -----
references_lines = [
    ("Schulhoff, S. et al. (2023). Ignore This Title and HackAPrompt: Exposing Systemic Vulnerabilities of LLMs through a Global Scale Prompt Hacking Competition. EMNLP 2023. arXiv:2311.16119.", False),
    ("Toyer, S. et al. (2024). Tensor Trust: Interpretable Prompt Injection Attacks from an Online Game. ICLR 2024. arXiv:2311.01011.", False),
    ("Yi, J. et al. (2025). Benchmarking and Defending Against Indirect Prompt Injection Attacks on Large Language Models (BIPIA). KDD 2025. arXiv:2312.14197.", False),
    ("Zhan, Q. et al. (2024). InjecAgent: Benchmarking Indirect Prompt Injections in Tool-Integrated Large Language Model Agents. ACL Findings 2024. arXiv:2403.02691.", False),
    ("Debenedetti, E. et al. (2024). AgentDojo: A Dynamic Environment to Evaluate Prompt Injection Attacks and Defenses for LLM Agents. NeurIPS 2024. arXiv:2406.13352.", False),
    ("Chen, S., Piet, J., Sitawarin, C., Wagner, D. (2025). StruQ: Defending Against Prompt Injection with Structured Queries. USENIX Security 2025. arXiv:2402.06363.", False),
    ("Hines, K. et al. (2024). Defending Against Indirect Prompt Injection Attacks With Spotlighting. Microsoft Research. arXiv:2403.14720.", False),
    ("Hung, K.-H. et al. (2025). Attention Tracker: Detecting Prompt Injection Attacks in LLMs. NAACL Findings 2025. arXiv:2411.00348.", False),
    ("Hu, T. et al. (2025). Defending against Indirect Prompt Injection by Instruction Detection. arXiv:2505.06311.", False),
    ("OWASP Foundation (2025). LLM01:2025 Prompt Injection. OWASP Top 10 for LLM Applications. https://genai.owasp.org/llmrisk/llm01-prompt-injection/.", False),
    ("Meta (2025). Llama-Prompt-Guard-2-86M. https://huggingface.co/meta-llama/Llama-Prompt-Guard-2-86M.", False),
    ("ProtectAI (2024). deberta-v3-base-prompt-injection-v2. https://huggingface.co/protectai/deberta-v3-base-prompt-injection-v2.", False),
]
last = append_body(headings["references"], references_lines)

doc.save(DST)
print("Saved to", DST)
